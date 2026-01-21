
import os
import time

from fastapi.testclient import TestClient
from docx import Document
from main import app, DB_FILE, OUTPUT_DIR

client = TestClient(app)

# Ensure clean state
if os.path.exists(DB_FILE):
    os.remove(DB_FILE)
    from main import init_db
    init_db()

def create_dummy_docx(filename):
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    doc.save(filename)

def test_upload_and_conversion():
    filename = "test_doc.docx"
    create_dummy_docx(filename)
    
    with open(filename, "rb") as f:
        response = client.post("/upload", files={"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    
    assert response.status_code == 200
    data = response.json()
    assert "job_id" in data
    job_id = data["job_id"]
    
    # Wait for background task (TestClient runs background tasks synchronously usually, but let's check)
    # Actually TestClient executes background tasks after the response is sent.
    
    # Poll for status
    for _ in range(10):
        response = client.get(f"/status/{job_id}")
        status_data = response.json()
        if status_data["status"] == "completed":
            break
        time.sleep(0.5)
        
    assert status_data["status"] == "completed"
    
    # Download
    response = client.get(f"/download/{job_id}")
    assert response.status_code == 200
    content = response.content.decode("utf-8")
    # Download
    response = client.get(f"/download/{job_id}")
    assert response.status_code == 200
    content = response.content.decode("utf-8")
    
    # Check for GIFT format markers
    assert "::" in content
    assert "{" in content
    assert "=" in content
    assert "}" in content
    
    # Cleanup
    os.remove(filename)

if __name__ == "__main__":
    try:
        test_upload_and_conversion()
        print("Test Passed!")
    except Exception as e:
        print(f"Test Failed: {e}")
        import traceback
        traceback.print_exc()
